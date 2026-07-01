"""Render ScubaDrift docs to a single .docx with house diagrams + illustrative art.

- Authors the blueprint SVGs (figures.py) and illustrative SVGs (illustrative.py).
- Rasterizes each SVG -> PNG at 2x with cairosvg (ADR-093: PNG is the portable
  artifact every doc format embeds).
- Converts the Markdown docs (headings, tables, code, lists, quotes, inline
  formatting) into a styled Word document, inserting figures at named anchors.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

import cairosvg
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import figures
import illustrative

HERE = Path(__file__).resolve().parent
PKG = HERE.parent  # tools/scubadrift (build/ lives inside the package)
DOCS = PKG / "docs"
FIG_SVG = PKG / "docs" / "figures"
FIG_PNG = HERE / "png"
OUT = PKG / "ScubaDrift-Documentation.docx"

NAVY = RGBColor(0x0D, 0x1B, 0x2E)
TEAL = RGBColor(0x1E, 0x8C, 0x8C)
BODYC = RGBColor(0x33, 0x41, 0x4F)
MUTED = RGBColor(0x6B, 0x7A, 0x8D)
CODEBG = "F2F4F7"
HEADER_FONT = "Georgia"
BODY_FONT = "Arial"
MONO_FONT = "Consolas"

# Which docs, in reading order. Consolidated into one complete source file so
# the whole documentation set builds from a single Markdown input.
DOC_ORDER = ["documentation.md"]

# Figure placement: doc -> list of (heading-substring anchor, png filename, where).
# "before" inserts the figure just before the matched heading (opener); "after"
# inserts right after it; "top" inserts before the first heading.
FIGURES: dict[str, list[tuple[str, str, str]]] = {
    "documentation.md": [
        ("Architecture", "scubadrift-diagram-01-architecture.png", "after"),
        ("The three dispositions", "scubadrift-image-01-signal-in-noise.png", "before"),
        ("The three dispositions", "scubadrift-diagram-02-disposition-gate.png", "after"),
        ("Run-to-run drift", "scubadrift-diagram-03-drift-taxonomy.png", "after"),
        ("Risk-acceptance lifecycle", "scubadrift-image-02-expiry-horizon.png", "before"),
        ("Risk-acceptance lifecycle", "scubadrift-diagram-04-exception-lifecycle.png", "after"),
        ("Gating a pipeline", "scubadrift-diagram-05-conmon-gate.png", "after"),
        ("Continuous monitoring", "scubadrift-diagram-06-conmon-loop.png", "after"),
    ],
}

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def rasterize(svgs: list[Path]) -> None:
    FIG_PNG.mkdir(parents=True, exist_ok=True)
    for svg in svgs:
        png = FIG_PNG / (svg.stem + ".png")
        cairosvg.svg2png(url=str(svg), write_to=str(png), scale=2.0)


def _shade(el, hexfill: str) -> None:
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexfill)
    pr.append(shd)


def add_inline(p, text: str, size: int = 11) -> None:
    for part in re.split(INLINE, text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = MONO_FONT
            r.font.size = Pt(size - 1.5)
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            r = p.add_run(m.group(1) if m else part)
        else:
            r = p.add_run(part)
        if not r.font.name:
            r.font.name = BODY_FONT


def add_heading(doc, text: str, level: int) -> None:
    h = doc.add_heading("", level=level)
    r = h.add_run(text)
    r.font.name = HEADER_FONT
    r.font.color.rgb = NAVY if level <= 2 else RGBColor(0x1E, 0x3A, 0x66)
    r.bold = True


def add_code(doc, lines: list[str]) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = True
    cell = tbl.cell(0, 0)
    _shade(cell._tc, CODEBG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for i, ln in enumerate(lines):
        r = p.add_run(ln)
        r.font.name = MONO_FONT
        r.font.size = Pt(8.5)
        if i < len(lines) - 1:
            r.add_break()


def add_table(doc, rows: list[list[str]]) -> None:
    header, *body = rows
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for j, cell_text in enumerate(header):
        c = t.rows[0].cells[j]
        _shade(c._tc, "0D1B2E")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = c.paragraphs[0]
        add_inline(para, cell_text, size=10)
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.bold = True
            run.font.size = Pt(10)
    for row in body:
        cells = t.add_row().cells
        for j, cell_text in enumerate(row):
            if j >= len(cells):
                break
            para = cells[j].paragraphs[0]
            add_inline(para, cell_text, size=10)
            for run in para.runs:
                if run.font.size is None:
                    run.font.size = Pt(10)


def add_figure(doc, png_name: str) -> None:
    png = FIG_PNG / png_name
    if not png.exists():
        print(f"  WARN missing figure {png_name}", file=sys.stderr)
        return
    doc.add_picture(str(png), width=Inches(6.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(6)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(10)


def _is_table_row(line: str) -> bool:
    return line.lstrip().startswith("|") and line.rstrip().endswith("|")


def _is_sep_row(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|(\s*:?-+:?\s*\|)+\s*", line))


def _cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def render_markdown(doc, md: str, doc_figs: list[tuple[str, str, str]]) -> None:
    lines = md.splitlines()
    i = 0
    placed_top = False
    # place a "top" opener figure before the first heading
    top_fig = next((f for f in doc_figs if f[2] == "top"), None)
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code
        if stripped.startswith("```"):
            j = i + 1
            code: list[str] = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                code.append(lines[j])
                j += 1
            add_code(doc, code if code else [""])
            i = j + 1
            continue

        # table
        if _is_table_row(line):
            rows = []
            j = i
            while j < len(lines) and _is_table_row(lines[j]):
                if not _is_sep_row(lines[j]):
                    rows.append(_cells(lines[j]))
                j += 1
            if rows:
                add_table(doc, rows)
            i = j
            continue

        # heading
        m = re.match(r"(#{1,6})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # opener figure before the first heading
            if not placed_top and top_fig is not None:
                add_figure(doc, top_fig[1])
                placed_top = True
            # "before" figures for this heading
            for anchor, png, where in doc_figs:
                if where == "before" and anchor in text:
                    add_figure(doc, png)
            add_heading(doc, re.sub(r"[`*]", "", text), min(level, 4))
            # "after" figures for this heading
            for anchor, png, where in doc_figs:
                if where == "after" and anchor in text:
                    add_figure(doc, png)
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            buf = []
            j = i
            while j < len(lines) and lines[j].strip().startswith(">"):
                buf.append(lines[j].strip().lstrip(">").strip())
                j += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_inline(p, " ".join(buf))
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = MUTED
            i = j
            continue

        # list item (unordered or ordered), one nesting level via indent
        lm = re.match(r"(\s*)([-*]|\d+\.)\s+(.*)", line)
        if lm:
            indent = len(lm.group(1))
            ordered = bool(re.match(r"\d+\.", lm.group(2)))
            style = ("List Number" if ordered else "List Bullet") + (" 2" if indent >= 2 else "")
            try:
                p = doc.add_paragraph(style=style)
            except KeyError:
                p = doc.add_paragraph(style="List Bullet")
            add_inline(p, lm.group(3))
            i += 1
            continue

        # blank
        if not stripped:
            i += 1
            continue

        # paragraph (gather until blank / block start)
        buf = [line]
        j = i + 1
        while j < len(lines):
            nxt = lines[j]
            if (
                not nxt.strip()
                or nxt.strip().startswith(("#", ">", "```", "|"))
                or re.match(r"(\s*)([-*]|\d+\.)\s+", nxt)
            ):
                break
            buf.append(nxt)
            j += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(s.strip() for s in buf))
        i = j


def style_base(doc) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODYC


def title_page(doc) -> None:
    add_figure(doc, "scubadrift-image-00-cover.png")
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("ScubaDrift")
    r.font.name = HEADER_FONT
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run("Drift, exception-lifecycle, and conmon gating on top of CISA ScubaGear")
    rs.font.name = BODY_FONT
    rs.font.size = Pt(13)
    rs.font.color.rgb = MUTED
    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rv = ver.add_run("Version 0.1.0  ·  Complete documentation")
    rv.font.name = BODY_FONT
    rv.font.size = Pt(11)
    rv.font.color.rgb = TEAL
    doc.add_page_break()


def main() -> int:
    # 1. author SVGs
    figures.build_all(FIG_SVG)
    illustrative.build_all(FIG_SVG)
    svgs = sorted(FIG_SVG.glob("scubadrift-*.svg"))
    print(f"authored {len(svgs)} SVG figures")

    # 2. rasterize at 2x
    rasterize(svgs)
    print(f"rasterized {len(list(FIG_PNG.glob('*.png')))} PNGs")

    # 3. assemble docx
    doc = Document()
    style_base(doc)
    title_page(doc)
    for k, name in enumerate(DOC_ORDER):
        md = (DOCS / name).read_text(encoding="utf-8")
        render_markdown(doc, md, FIGURES.get(name, []))
        if k < len(DOC_ORDER) - 1:
            doc.add_page_break()
    doc.save(str(OUT))
    print(f"wrote {OUT}  ({OUT.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
