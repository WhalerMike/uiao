"""Generate rich Word documents from UIAO canon data using docxtpl.

Produces publication-quality DOCX with native Word styles, auto-TOC,
headers/footers with classification markings, embedded images, and
properly formatted compliance tables.

Usage:
    python scripts/generate_rich_docx.py

Requires:
    pip install docxtpl python-docx pyyaml

If a Word template exists at templates/leadership_briefing_v1.0.docx,
it is used as the base (with {{jinja}} placeholders). Otherwise, the
script programmatically builds a styled document from scratch.
"""
import yaml
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

try:
    from docxtpl import DocxTemplate
    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parent.parent
CANON = ROOT / "canon" / "uiao_leadership_briefing_v1.0.yaml"
DATA_DIR = ROOT / "data"
TEMPLATES_DIR = ROOT / "templates"
VISUALS_DIR = ROOT / "visuals"
EXPORTS_DIR = ROOT / "exports"


def load_context():
    """Load canon + data files, same pattern as generate_docs.py."""
    context = {}
    if DATA_DIR.exists():
        for yml_file in sorted(DATA_DIR.glob("*.yml")):
            key = yml_file.stem.replace("-", "_")
            with yml_file.open("r", encoding="utf-8") as f:
                content = yaml.safe_load(f)
            if content:
                if isinstance(content, dict):
                    context.update(content)
                context[key] = content
    if CANON.exists():
        with CANON.open("r", encoding="utf-8") as f:
            canon = yaml.safe_load(f)
        if canon:
            context.update(canon)
    return context


def _add_classification_header(doc, classification="CUI"):
    """Add header with classification marking to all sections."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = classification
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        run.font.bold = True
        # Footer
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = f"{classification} | UIAO Program | Generated {datetime.now():%Y-%m-%d}"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.runs[0]
        fr.font.size = Pt(7)
        fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_heading(doc, text, level=1):
    """Add a heading with consistent styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Navy
    return h


def _add_narrative(doc, text):
    """Add a narrative paragraph with proper formatting."""
    if not text:
        return
    for para_text in str(text).split("\n\n"):
        p = doc.add_paragraph(para_text.strip())
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.size = Pt(11)


def _add_image_safe(doc, image_name, width=Inches(5.5)):
    """Add an image if it exists, skip gracefully if not."""
    img_path = VISUALS_DIR / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    logger.warning("Image not found: %s", img_path)
    return False


def _add_compliance_table(doc, matrix):
    """Add the Unified Compliance Matrix as a formatted Word table."""
    headers = ["UIAO Pillar", "CISA ZT Pillar", "Target Maturity",
               "NIST 800-53 Controls", "Mission Impact"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for entry in matrix:
        row = table.add_row()
        row.cells[0].text = entry.get("pillar", "")
        row.cells[1].text = entry.get("cisa_pillar", "")
        row.cells[2].text = entry.get("cisa_maturity", "")
        controls = entry.get("nist_controls", [])
        row.cells[3].text = ", ".join(controls) if isinstance(controls, list) else str(controls)
        row.cells[4].text = entry.get("impact_statement", "")
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def _add_evidence_table(doc):
    """Add the FedRAMP Audit Evidence Summary table."""
    evidence_map = [
        ("V1: Identity-to-IP Mapping", "U + A (The Gate)", "IA-2, AC-19, CM-8",
         "uiao-vibrant-u-plus-a-mapping.png"),
        ("V2: INR Fabric", "O (The Network)", "AC-4",
         "uiao-vibrant-o-pillar-inr-fabric.png"),
        ("V3: 20x Governance Loop", "Governance (The Hub)", "CA-7, IR-4",
         "uiao-vibrant-20x-governance-hub.png"),
        ("V4: Modernization Atlas", "Strategy (The Journey)", "Program Vision / TIC 3.0",
         "uiao-vibrant-modernization-atlas.png"),
        ("V5: Cryptographic Trust Chain", "Security (The Lock)", "SC-8",
         "uiao-vibrant-cryptographic-trust-chain.png"),
    ]
    headers = ["Visual Title", "Architectural Pillar", "NIST Control(s)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for title, pillar, controls, _img in evidence_map:
        row = table.add_row()
        row.cells[0].text = title
        row.cells[1].text = pillar
        row.cells[2].text = controls
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


def build_from_scratch(context):
    """Build a complete styled DOCX programmatically."""
    doc = Document()
    lb = context.get("leadership_briefing", {})
    if not isinstance(lb, dict):
        lb = {}
    classification = context.get("classification", "CUI")

    _add_classification_header(doc, classification)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Unified Identity-Addressing-Overlay Architecture (UIAO)")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"Leadership Briefing - Version {context.get('version', '1.0')}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Generated: {datetime.now():%B %d, %Y}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_page_break()

    # TOC placeholder (Word updates on open)
    _add_heading(doc, "Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    toc_p.add_run("[Table of Contents - right-click and select 'Update Field' in Word]")
    toc_p.runs[0].font.size = Pt(9)
    toc_p.runs[0].font.italic = True
    doc.add_page_break()

    # Executive Summary
    _add_heading(doc, "Executive Summary")
    _add_narrative(doc, lb.get("executive_summary"))

    # Program Overview
    _add_heading(doc, "Program Overview")
    _add_narrative(doc, lb.get("program_overview"))

    # Modernization Need
    _add_heading(doc, "Why Modernization Is Required")
    _add_narrative(doc, lb.get("modernization_need"))

    # Program Vision
    _add_heading(doc, "Program Vision")
    _add_narrative(doc, lb.get("program_vision"))

    # Five Control Planes
    _add_heading(doc, "The Five Control Planes")
    for i, plane in enumerate(lb.get("control_planes", []), 1):
        if isinstance(plane, dict):
            _add_heading(doc, f"{i}. {plane.get('name', '')}", level=2)
            _add_narrative(doc, plane.get("narrative"))

    # Seven Core Concepts
    _add_heading(doc, "Seven Core Concepts")
    for i, concept in enumerate(lb.get("core_concepts", []), 1):
        if isinstance(concept, dict):
            _add_heading(doc, f"{i}. {concept.get('name', '')}", level=2)
            _add_narrative(doc, concept.get("narrative"))

    # Frozen State
    _add_heading(doc, "Frozen State Analysis")
    _add_narrative(doc, lb.get("frozen_state"))

    # Outcomes
    _add_heading(doc, "Program Outcomes")
    _add_narrative(doc, lb.get("outcomes"))

    # Vibrant Visualizations
    doc.add_page_break()
    _add_heading(doc, "Vibrant Visualizations")
    visuals = [
        ("Modernization Journey", "uiao-vibrant-modernization-atlas.png"),
        ("FedRAMP 20x Governance Loop", "uiao-vibrant-20x-governance-hub.png"),
        ("Identity-to-IP Architecture", "uiao-vibrant-u-plus-a-mapping.png"),
    ]
    for title_text, img in visuals:
        _add_heading(doc, title_text, level=2)
        _add_image_safe(doc, img)

    # Maturity radar chart (if generated)
    radar_path = VISUALS_DIR / "dynamic-maturity-radar.png"
    if radar_path.exists():
        doc.add_page_break()
        _add_heading(doc, "CISA Zero Trust Maturity Assessment")
        _add_image_safe(doc, "dynamic-maturity-radar.png")

    # FedRAMP Evidence Summary
    doc.add_page_break()
    _add_heading(doc, "FedRAMP 20x Audit Evidence Summary")
    _add_narrative(doc, "Direct mapping of UIAO architecture to NIST 800-53 Rev 5 controls.")
    _add_evidence_table(doc)

    # Compliance Matrix
    doc.add_page_break()
    _add_heading(doc, "Unified Compliance & Maturity Matrix")
    matrix = context.get("unified_compliance_matrix", [])
    if matrix:
        _add_compliance_table(doc, matrix)
    p = doc.add_paragraph()
    p.add_run(
        "Auditor Note: All controls listed above are continuously monitored "
        "via the UIAO Governance Plane (V3) and reported through the "
        "ServiceNow SCuBA integration."
    ).font.italic = True

    return doc


def main():
    logger.info("Loading UIAO context...")
    context = load_context()

    docx_dir = EXPORTS_DIR / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    out_path = docx_dir / "UIAO_Leadership_Briefing_v1.0.docx"

    # Try template-based approach first
    tpl_path = TEMPLATES_DIR / "leadership_briefing_v1.0.docx"
    if HAS_DOCXTPL and tpl_path.exists():
        logger.info("Using docxtpl template: %s", tpl_path)
        tpl = DocxTemplate(str(tpl_path))
        context["today"] = datetime.now().strftime("%B %d, %Y")
        context["compliance_table"] = context.get("unified_compliance_matrix", [])
        tpl.render(context)
        tpl.save(str(out_path))
    else:
        logger.info("Building styled DOCX from scratch (no .docx template found)...")
        doc = build_from_scratch(context)
        doc.save(str(out_path))

    logger.info("Rich DOCX exported -> %s", out_path)
    logger.info("  Sections: Title, TOC, Exec Summary, 5 Control Planes,")
    logger.info("            7 Core Concepts, Visualizations, Evidence, Compliance Matrix")
    logger.info("  Open in Word and press Ctrl+A then F9 to update TOC fields.")


if __name__ == "__main__":
    main()
