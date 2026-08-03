#!/usr/bin/env python3
"""Concatenate per-page .docx renders into one bundle.docx per section.

Walks ``_site/customer-documents/<section>/`` (the directory the
quarto.yml `assemble` job populates), finds every page-level .docx,
sorts them deterministically by path, and uses ``docxcompose`` to
concatenate them into a single ``<section>-bundle.docx`` written next
to the section's index.html / index.docx.

Customers get a downloadable "complete pack" per section without
having to click through every individual page. The per-page .docx
files remain untouched alongside the bundle.

Why docxcompose vs. raw Pandoc
------------------------------
Pandoc can concatenate .qmd source, but Quarto-specific shortcodes
(``{{< include >}}``) don't process through raw Pandoc, image
``--resource-path`` resolution is brittle across the nested
customer-documents tree, and Pandoc-emitted DOCX loses some styling
nuances vs. Quarto's. ``docxcompose`` operates on the already-rendered
.docx outputs Quarto produced and preserves their styles, numbered
lists, embedded images, and headers/footers verbatim.

Skipped from bundles
--------------------
- Files whose .docx wasn't produced (Quarto's matrix skips empty stubs).
- The section's own ``index.docx`` (nav surface, not content).
- ``ROADMAP.docx``, ``document-index.docx`` (cross-cutting meta surfaces).
- ``<section>-bundle.docx`` itself (don't recursively bundle past runs).
- Pre-existing ``.docx`` artifacts under nested ``images/`` directories.

Stripped from each chapter before concatenation
-----------------------------------------------
- Pandoc's auto-emitted ``Table of contents`` block (a ``<w:sdt>``
  Structured Document Tag with ``<w:docPartGallery w:val="Table of
  Contents"/>``). Per-chapter docs need TOCs for standalone download
  but 16 mini-TOCs interleaved through a bundle is noise. Stripping
  happens in memory on each loaded chapter; the on-disk per-chapter
  ``.docx`` files are untouched.

Usage
-----
::

    python scripts/bundle_section_docx.py \\
        --site-root _site/customer-documents \\
        --section whitepapers

    # Bundle every top-level section:
    python scripts/bundle_section_docx.py \\
        --site-root _site/customer-documents \\
        --all
"""

from __future__ import annotations

import argparse
import re
import shutil
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

# Imported lazily inside main() so --help works even without the dep
# installed (useful for CI failure diagnostics).


# Top-level sections to bundle. Order is arbitrary — the bundler runs
# them independently. Keeping the list explicit (rather than scanning
# the directory) so a typo in the site layout doesn't silently produce
# zero bundles, and so a new top-level section is an intentional add.
DEFAULT_SECTIONS = [
    "adapter-specs",
    "architecture-series",
    "case-studies",
    "executive-briefs",
    "executive-governance-series",
    "orgcomp-series",
    "modernization-specs",
    "operational-guides",
    # Nested sub-section: a focused bundle of just the OrgPath Intune/Arc
    # implementation guides, in addition to the whole operational-guides
    # bundle (which still includes these pages via its recursive walk).
    # Resolved relative to the site root; the bundle filename uses only the
    # final path component (see _bundle_filename).
    "operational-guides/orgpath-implementation",
    # Additional nested operational-guides sub-sections that hold their own
    # multi-page content (in addition to the whole operational-guides bundle,
    # which still includes these pages via its recursive walk). Index-only
    # navigation stubs (directory-migration, platform-substrate,
    # transformation-engine, zero-trust-assessment) are intentionally omitted:
    # they have no local content .docx to bundle, so a bundle link would 404.
    "operational-guides/access-plane",
    "operational-guides/client-server-to-hybrid-cloud",
    "operational-guides/identity-orgtree",
    "operational-guides/network-transformation",
    "operational-guides/program-management",
    "operational-guides/target-surface",
    "operational-guides/uiao-modernization-program",
    "orgpath-narrative",
    "sql-server-narrative",
    # The reference-free twin of sql-server-narrative (same ten-book reading
    # sequence with internal cross-references stripped). Same chaptered Book_NN
    # + Book_NN_CPT_MM layout, so it gets both a whole-section bundle here and
    # per-book bundles via BOOK_BUNDLE_SECTIONS below — mirroring its twin.
    "sql-server-narrative-standalone",
    # The executable companion to sql-server-narrative. Single-file books
    # (Book_NN.qmd, no chapter splits), so it gets a whole-companion bundle
    # only — not per-book bundles (a single-page per-book bundle would just
    # duplicate the page's own Book_NN.docx render).
    "sql-server-implementation",
    "reference-architecture",
    "platform",
    "substrate",
    "validation-suites",
    "whitepapers",
    "ssot-network-physics",
]

# Filename stems (without .docx) skipped from every bundle.
SKIP_STEMS = {"index", "ROADMAP", "document-index", "TREE"}

# Sections whose "bundle" is a zip of the individual page .docx files rather
# than one docxcompose-merged .docx. whitepapers stays in DEFAULT_SECTIONS
# above (so --all still processes it) but is redirected to _zip_pages in
# main() instead of _bundle_one, because a single merged
# whitepapers-bundle.docx concatenating all ~22 long-form papers had grown
# too large to be a usable download; a zip keeps each paper as its own file
# (openable individually, no 20+-chapter Word doc to scroll through) at a
# fraction of the total size. Add a section name here to switch it from
# merged-docx to zipped-pages; see _zip_pages / _bundle_filename_zip.
ZIP_INSTEAD_OF_DOCX_SECTIONS = {"whitepapers"}

# Sections that are organized as multi-chapter "books" (Book_NN.qmd +
# Book_NN_CPT_MM.qmd). For these, in addition to the whole-section
# bundle, emit one Book_NN-bundle.docx per book so a reader can download
# a single book as one Word file. Keyed by section directory name.
BOOK_BUNDLE_SECTIONS = {"orgpath-narrative", "sql-server-narrative", "sql-server-narrative-standalone"}

# Sections organized as roman-numbered volumes of books (Vol_<N>_Book_NN_*).
# For these, in addition to the whole-section bundle, emit one
# Vol_<N>-<Theme>-Bundle.docx per volume — the whole-series bundle spans ~90+
# page documents and is unwieldy as a single Word file; a per-volume download
# is the unit a reader actually works with.
VOLUME_BUNDLE_SECTIONS = {"orgcomp-series"}

# Volume themes, keyed by roman volume token. Carried in the bundle filename
# ("Vol_I-Foundation & Transport-Bundle.docx") and the stamped running header
# so a downloaded file identifies its subject without opening it. Must match
# the Series-structure tables on docs/download/index.qmd and
# docs/customer-documents/orgcomp-series/index.qmd. Vol X's "(optional
# binding)" table annotation is deliberately not part of the theme.
VOLUME_THEMES = {
    "0": "Executive Summary, Questionnaire & Control Crosswalk",
    "I": "Foundation & Transport",
    "II": "Data Platform",
    "III": "Security Operations",
    "IV": "Governance & Assurance",
    "V": "Training & Certification",
    "VI": "Implementation",
    "VII": "ServiceNow Automation",
    "VIII": "Multi-Cloud DDI",
    "IX": "Day-2 Operations",
    "X": "Governance-Substrate Integration",
}


def _volume_bundle_filename(vol_id: str) -> str:
    """Per-volume bundle filename: ``Vol_<N>-<Theme>-Bundle.docx``.

    A volume with no registered theme falls back to the themeless
    ``Vol_<N>-bundle.docx`` form so an unmapped volume still bundles.
    Either form is recognised by the case-insensitive bundle-skip filter.
    """
    theme = VOLUME_THEMES.get(vol_id)
    if theme is None:
        return f"Vol_{vol_id}-bundle.docx"
    return f"Vol_{vol_id}-{theme}-Bundle.docx"


# Leading volume identifier in a series page filename,
# e.g. "Vol_IX_Book_03_..." -> "IX". Top-level series books only — the
# training-program and boundary subtrees don't carry the Vol_ prefix.
VOL_ID_RE = re.compile(r"^Vol_([0IVX]+)_")

# Leading book identifier in a page filename, e.g. "Book_07a_CPT_03" ->
# "Book_07a" and "Book_07a" (the landing page) -> "Book_07a". Used to
# group a section's page .docx into per-book buckets.
BOOK_ID_RE = re.compile(r"^(Book_\d+[a-z]?)")

# OpenXML namespace for the WordprocessingML schema. Used by the TOC-strip
# pass to locate Pandoc's auto-emitted Table-of-Contents block. Kept as a
# module-level constant so both the XPath query and the attribute lookup
# (Clark-notation ``{ns}val``) reference the same string.
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _strip_toc_blocks(doc) -> int:
    """Remove Pandoc auto-TOC blocks from a loaded Document, in place.

    Pandoc renders the chapter-level TOC as a ``<w:sdt>`` (Structured
    Document Tag) whose ``<w:docPartGallery>`` element carries the
    ``w:val="Table of Contents"`` marker. Returns the number of TOC
    blocks removed (0 if the chapter had no TOC, e.g. a stub or a doc
    rendered with ``toc: false``).

    Why strip at bundle time, not at chapter render time: per-chapter
    ``.docx`` files are also offered as standalone downloads on the
    site, and TOCs are useful there. Stripping happens on the in-memory
    Document loaded by the bundler; the on-disk per-chapter file is
    never rewritten.
    """
    body = doc.element.body
    ns = {"w": W_NS}
    removed = 0
    for sdt in body.findall(".//w:sdt", ns):
        gallery = sdt.find(".//w:docPartGallery", ns)
        if gallery is not None and gallery.get(f"{{{W_NS}}}val") == "Table of Contents":
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)
                removed += 1
    return removed


# Running-header rewrite for per-book bundles.
#
# The section reference doc (`_reference.docx`) wires the running header as two
# STYLEREF fields: "Heading 1" on the left (intended: book title) and
# "Heading 2" on the right (intended: chapter). But chapter titles are *also*
# Heading 1, so STYLEREF "Heading 1" resolves to the chapter from page 2 on and
# the book identity disappears from the header. For a per-book bundle the book
# is fixed, so we replace the left field with a static "Book_NN — Title" run
# (which can no longer be clobbered by a chapter heading) and repoint the right
# field at Heading 1 so it shows the chapter. Result: the book title persists
# on every page.
_HDR_RPR = '<w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:color w:val="616161"/><w:sz w:val="18"/></w:rPr>'
_LEFT_STYLEREF_FIELD = re.compile(
    r"<w:r>" + re.escape(_HDR_RPR) + r'<w:fldChar w:fldCharType="begin"/></w:r>'
    r'.*?STYLEREF "Heading 1".*?'
    r"<w:r>" + re.escape(_HDR_RPR) + r'<w:fldChar w:fldCharType="end"/></w:r>',
    re.DOTALL,
)


def _xml_escape(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _first_heading1_text(docx_path: Path) -> str | None:
    """Return the text of the first Heading-1 paragraph in a .docx, or None.

    For a per-book bundle the first Heading 1 is the landing page's
    ``Book_NN — Title`` heading.
    """
    with zipfile.ZipFile(docx_path) as z:
        doc = ET.fromstring(z.read("word/document.xml"))
    for p in doc.iter(f"{{{W_NS}}}p"):
        pPr = p.find(f"{{{W_NS}}}pPr")
        ps = pPr.find(f"{{{W_NS}}}pStyle") if pPr is not None else None
        if ps is not None and ps.get(f"{{{W_NS}}}val") == "Heading1":
            return "".join(t.text or "" for t in p.iter(f"{{{W_NS}}}t")).strip()
    return None


def _rewrite_header_xml(xml: str, book_title: str) -> tuple[str, int]:
    """Stamp the book title into one header part's XML. Returns (xml, n_left)."""
    literal = "<w:r>" + _HDR_RPR + '<w:t xml:space="preserve">' + _xml_escape(book_title) + "</w:t></w:r>"
    xml, n_left = _LEFT_STYLEREF_FIELD.subn(literal, xml, count=1)
    # Right field: show the chapter (Heading 1) rather than the section heading.
    xml = re.sub(r'STYLEREF "Heading 2"', 'STYLEREF "Heading 1"', xml, count=1)
    return xml, n_left


def _stamp_book_header(bundle_path: Path, book_title: str) -> int:
    """Rewrite every running-header part in a bundle to show the book title.

    Returns the number of header parts changed. No-op (returns 0, original
    file left untouched) if no header carries the expected STYLEREF "Heading 1"
    field, so it is safe to call on a bundle built from a different reference
    doc.
    """
    tmp = bundle_path.with_suffix(".stamp.docx")
    changed = 0
    with zipfile.ZipFile(bundle_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if re.fullmatch(r"word/header\d*\.xml", item.filename):
                new, n_left = _rewrite_header_xml(data.decode("utf-8"), book_title)
                if n_left:
                    changed += 1
                data = new.encode("utf-8")
            zout.writestr(item, data)
    if changed:
        shutil.move(str(tmp), str(bundle_path))
    else:
        tmp.unlink()
    return changed


def _bundle_filename(section: str) -> str:
    """Return the bundle filename for a section path.

    Uses only the final path component so a nested section path (e.g.
    ``operational-guides/orgpath-implementation``) yields
    ``orgpath-implementation-bundle.docx`` rather than a name with an
    embedded directory separator. Top-level sections are unaffected
    (``operational-guides`` -> ``operational-guides-bundle.docx``).
    """
    return f"{Path(section).name}-bundle.docx"


def _bundle_filename_zip(section: str) -> str:
    """Return the zip-bundle filename for a section path (see _zip_pages)."""
    return f"{Path(section).name}-bundle.zip"


def _zip_pages(site_root: Path, section: str) -> tuple[bool, str]:
    """Zip a section's page .docx files individually instead of merging them.

    Used for sections in ZIP_INSTEAD_OF_DOCX_SECTIONS where one
    docxcompose-merged .docx would be an unwieldy single download. Reuses
    _collect_docx so the page set (and its index/ROADMAP/prior-bundle
    exclusions) matches exactly what _bundle_one would have merged — only
    the packaging differs. Each page is stored under its filename only
    (no directory prefix), matching the flat layout a reader expects when
    unzipping "the whitepapers" into one folder.
    """
    section_dir = site_root / section
    if not section_dir.is_dir():
        return False, f"{section}: directory not found at {section_dir}"

    zip_name = _bundle_filename_zip(section)
    zip_path = section_dir / zip_name

    pages = _collect_docx(section_dir, zip_name)
    if not pages:
        return False, f"{section}: no page .docx files found under {section_dir}"

    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as z:
        for page in pages:
            z.write(page, arcname=page.name)

    return (
        True,
        f"{section}: zipped {len(pages)} page(s) -> {zip_path.relative_to(site_root.parent)}",
    )


# Roman-numeral volume ranks for the OrgComp series reading order. Plain
# lexicographic sort misplaces IX before V and, worse, puts the
# OrgComp-Training-Program/ subtree FIRST (ASCII '-' < '_' < letters), so the
# series bundle used to open with ~28 academy pages before any book — reading
# like a training document rather than the series.
_ROMAN_RANK = {"0": 0, "I": 1, "II": 2, "III": 3, "IV": 4, "V": 5, "VI": 6, "VII": 7, "VIII": 8, "IX": 9, "X": 10}
_VOL_BOOK_RE = re.compile(r"^Vol_([0IVX]+)_Book_(\w+?)_")


def _orgcomp_series_key(section_dir: Path, docx: Path):
    """Reading-order sort key for the orgcomp-series bundle.

    Group 0: the Vol 0–X books, in volume order (roman-aware) then book
    number. Group 1: the boundary annexes. Group 2: top-level companion
    specs (Evidence Contract, etc.). Group 3: the Training Program — the
    academy belongs at the END of the series bundle, not the front.
    """
    rel = docx.relative_to(section_dir)
    top = rel.parts[0]
    m = _VOL_BOOK_RE.match(rel.name)
    if len(rel.parts) == 1 and m and m.group(1) in _ROMAN_RANK:
        return (0, _ROMAN_RANK[m.group(1)], m.group(2), str(rel))
    if top == "boundary":
        return (1, 0, "", str(rel))
    if top == "OrgComp-Training-Program":
        return (3, 0, "", str(rel))
    return (2, 0, "", str(rel))


def _collect_docx(section_dir: Path, bundle_name: str) -> list[Path]:
    """Return the sorted list of page .docx files to include in a bundle.

    Walks the section tree, applies the SKIP_STEMS filter, and excludes
    the bundle's own output file so a re-run doesn't fold the previous
    bundle into itself. The orgcomp-series section sorts in reading order
    (books → boundary → specs → training program) instead of raw path
    order — see _orgcomp_series_key.
    """
    out: list[Path] = []
    for docx in sorted(section_dir.rglob("*.docx")):
        # Skip the section bundle and any per-book/per-volume bundle so a
        # re-run (or the per-book/per-volume pass below) never folds a prior
        # bundle into itself. Case-insensitive: per-volume bundles end with
        # "-Bundle.docx" (capital B, from the themed filename form).
        if docx.name.lower().endswith("-bundle.docx"):
            continue
        if docx.stem in SKIP_STEMS:
            continue
        out.append(docx)
    if section_dir.name == "orgcomp-series":
        out.sort(key=lambda d: _orgcomp_series_key(section_dir, d))
    return out


def _bundle_one(site_root: Path, section: str) -> tuple[bool, str]:
    """Bundle one section. Returns ``(ok, message)``."""
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docxcompose.composer import Composer

    section_dir = site_root / section
    if not section_dir.is_dir():
        return False, f"{section}: directory not found at {section_dir}"

    bundle_name = _bundle_filename(section)
    bundle_path = section_dir / bundle_name

    pages = _collect_docx(section_dir, bundle_name)
    if not pages:
        return False, f"{section}: no page .docx files found under {section_dir}"

    # First page becomes the master — its styles, headers, and footers
    # carry through to the bundle. Subsequent pages append in order, each
    # preceded by a page break so individual chapters don't flow into one
    # another (the default docxcompose behavior is back-to-back concat
    # with no separator).
    master = Document(str(pages[0]))
    tocs_stripped = _strip_toc_blocks(master)
    composer = Composer(master)
    for page in pages[1:]:
        # Insert an empty paragraph with a page break at the end of the
        # current master body, then append the next doc's content. The
        # break sits between docs, not inside either one, so neither
        # source doc is mutated and styles/headers stay intact.
        master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        chapter = Document(str(page))
        tocs_stripped += _strip_toc_blocks(chapter)
        composer.append(chapter)
    composer.save(str(bundle_path))

    return (
        True,
        f"{section}: bundled {len(pages)} pages "
        f"(stripped {tocs_stripped} per-chapter TOC block(s)) -> "
        f"{bundle_path.relative_to(site_root.parent)}",
    )


def _book_sort_key(book_id: str) -> tuple[int, int]:
    """Order book ids so e.g. Book_07 precedes Book_07a precedes Book_08."""
    m = re.match(r"Book_(\d+)([a-z]?)$", book_id)
    if not m:
        return (0, 0)
    return (int(m.group(1)), 1 if m.group(2) else 0)


def _bundle_books(site_root: Path, section: str) -> list[tuple[bool, str]]:
    """Emit one ``Book_NN-bundle.docx`` per book within a chaptered section.

    Groups the section's page .docx by the leading ``Book_NN[a]`` filename
    token, concatenates each book's pages in filename order (landing page
    first, then chapters), and writes ``Book_NN-bundle.docx`` beside the
    book's landing page. Mirrors :func:`_bundle_one`'s TOC-strip and
    page-break behavior. Returns one ``(ok, message)`` tuple per book.
    """
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docxcompose.composer import Composer

    section_dir = site_root / section
    if not section_dir.is_dir():
        return [(False, f"{section}: directory not found at {section_dir}")]

    # Bucket page docx by book id. Files directly named Book_NN*.docx
    # qualify; section/per-book bundles and skip-stems are excluded.
    books: dict[str, list[Path]] = {}
    for docx in sorted(section_dir.rglob("*.docx")):
        if docx.name.lower().endswith("-bundle.docx"):
            continue
        if docx.stem in SKIP_STEMS:
            continue
        m = BOOK_ID_RE.match(docx.stem)
        if not m:
            continue
        books.setdefault(m.group(1), []).append(docx)

    results: list[tuple[bool, str]] = []
    for book_id in sorted(books, key=_book_sort_key):
        pages = books[book_id]
        bundle_path = section_dir / f"{book_id}-bundle.docx"

        master = Document(str(pages[0]))
        tocs_stripped = _strip_toc_blocks(master)
        composer = Composer(master)
        for page in pages[1:]:
            master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            chapter = Document(str(page))
            tocs_stripped += _strip_toc_blocks(chapter)
            composer.append(chapter)
        composer.save(str(bundle_path))

        # Stamp the book title into the running header so "Book_NN — Title"
        # persists on every page instead of being clobbered by chapter
        # headings. Falls back to the book id if no landing Heading 1 exists.
        book_title = _first_heading1_text(bundle_path) or book_id
        stamped = _stamp_book_header(bundle_path, book_title)

        results.append(
            (
                True,
                f"{section}/{book_id}: bundled {len(pages)} page(s) "
                f"(stripped {tocs_stripped} per-chapter TOC block(s); "
                f"header -> {book_title!r} on {stamped} part(s)) -> "
                f"{bundle_path.relative_to(site_root.parent)}",
            )
        )

    if not results:
        results.append((False, f"{section}: no Book_NN page .docx found for per-book bundles"))
    return results


def _bundle_volumes(site_root: Path, section: str) -> list[tuple[bool, str]]:
    """Emit one ``Vol_<N>-<Theme>-Bundle.docx`` per volume of a volume-structured section.

    Groups the section's TOP-LEVEL ``Vol_<N>_Book_NN_*.docx`` pages by their
    roman volume token, concatenates each volume's books in filename order
    (Book_00, Book_00a, Book_01, ...), and writes the themed bundle filename
    (see :func:`_volume_bundle_filename`) beside them. Subtrees (training program, boundary annexes) don't carry
    the ``Vol_`` prefix and are deliberately excluded — a volume bundle is
    the volume's books, nothing else. Mirrors :func:`_bundle_books`'s
    TOC-strip, page-break, and header-stamp behavior.
    """
    from docx import Document
    from docx.enum.text import WD_BREAK
    from docxcompose.composer import Composer

    section_dir = site_root / section
    if not section_dir.is_dir():
        return [(False, f"{section}: directory not found at {section_dir}")]

    volumes: dict[str, list[Path]] = {}
    for docx in sorted(section_dir.glob("*.docx")):
        if docx.name.lower().endswith("-bundle.docx"):
            continue
        if docx.stem in SKIP_STEMS:
            continue
        m = VOL_ID_RE.match(docx.stem)
        if not m or m.group(1) not in _ROMAN_RANK:
            continue
        volumes.setdefault(m.group(1), []).append(docx)

    results: list[tuple[bool, str]] = []
    for vol_id in sorted(volumes, key=_ROMAN_RANK.get):
        pages = volumes[vol_id]
        bundle_path = section_dir / _volume_bundle_filename(vol_id)

        master = Document(str(pages[0]))
        tocs_stripped = _strip_toc_blocks(master)
        composer = Composer(master)
        for page in pages[1:]:
            master.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            chapter = Document(str(page))
            tocs_stripped += _strip_toc_blocks(chapter)
            composer.append(chapter)
        composer.save(str(bundle_path))

        vol_label = "Volume 0" if vol_id == "0" else f"Volume {vol_id}"
        theme = VOLUME_THEMES.get(vol_id)
        title = f"Federal Organization Compliance — {vol_label}"
        if theme:
            title = f"{title} — {theme}"
        stamped = _stamp_book_header(bundle_path, title)

        results.append(
            (
                True,
                f"{section}/Vol_{vol_id}: bundled {len(pages)} book(s) "
                f"(stripped {tocs_stripped} per-book TOC block(s); "
                f"stamped {stamped} header part(s)) -> {bundle_path.name}",
            )
        )
    if not results:
        results.append((False, f"{section}: no Vol_<N> page .docx found for per-volume bundles"))
    return results


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n\n", 1)[0])
    parser.add_argument(
        "--site-root",
        type=Path,
        required=True,
        help="Path to the assembled customer-documents directory "
        "(typically `_site/customer-documents` inside the Quarto deploy).",
    )
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--section", help="Bundle a single section by name.")
    group.add_argument("--all", action="store_true", help="Bundle every default section.")
    parser.add_argument(
        "--canon-modernization",
        action="store_true",
        help="DEPRECATED no-op (kept for caller-flag compatibility). Per ADR-083, "
        "canon modernization pages moved into customer-documents at "
        "reference-architecture/, so they are now bundled by --all as a regular "
        "section. The flag will be removed in a follow-up PR once the Quarto "
        "workflow drops it from its invocation.",
    )
    args = parser.parse_args(argv)

    site_root = args.site_root.resolve()
    if not site_root.is_dir():
        print(f"error: --site-root not a directory: {site_root}", file=sys.stderr)
        return 1

    sections = [args.section] if args.section else DEFAULT_SECTIONS

    failures = 0
    for section in sections:
        if section in ZIP_INSTEAD_OF_DOCX_SECTIONS:
            ok, msg = _zip_pages(site_root, section)
        else:
            ok, msg = _bundle_one(site_root, section)
        prefix = "OK   " if ok else "WARN "
        print(f"{prefix} {msg}")
        if not ok:
            # Missing-section warnings shouldn't fail the run if --all is
            # used — a section may legitimately have no rendered .docx
            # (e.g., only stubs). Treat as non-fatal but report.
            failures += 1

        # Chaptered sections also get one Book_NN-bundle.docx per book.
        # Per-book failures are logged but never fatal — a missing book
        # bundle must not block the site deploy.
        if section in BOOK_BUNDLE_SECTIONS:
            for ok_b, msg_b in _bundle_books(site_root, section):
                print(f"{'OK   ' if ok_b else 'WARN '} {msg_b}")

        # Volume-structured sections also get one Vol_<N>-bundle.docx per
        # volume — same non-fatal treatment as the per-book pass.
        if section in VOLUME_BUNDLE_SECTIONS:
            for ok_v, msg_v in _bundle_volumes(site_root, section):
                print(f"{'OK   ' if ok_v else 'WARN '} {msg_v}")

    # --canon-modernization is a deprecated no-op per ADR-083 (canon
    # modernization pages now live at customer-documents/reference-architecture/
    # and are bundled by --all as a regular section). Log if the flag was passed
    # so the workflow author notices.
    if args.canon_modernization:
        print(
            "INFO  --canon-modernization is a no-op per ADR-083; "
            "reference-architecture is bundled by --all as a regular section."
        )

    if args.section and failures:
        # Single-section mode: caller asked for a specific bundle, so
        # any failure is fatal.
        return 1
    return 0


if __name__ == "__main__":
    sys.exit(main())
