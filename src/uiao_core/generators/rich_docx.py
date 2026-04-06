"""Rich DOCX generator for UIAO leadership briefing.

Migrated from scripts/generate_rich_docx.py into the uiao_core package.
Produces publication-quality DOCX with native Word styles, auto-TOC,
headers/footers with classification markings, embedded images, and
properly formatted compliance tables.

References: ADR-0004
"""

# isort: skip_file
from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

try:
    from docxtpl import DocxTemplate

    HAS_DOCXTPL = True
except ImportError:
    HAS_DOCXTPL = False

from uiao_core.utils.context import get_settings, load_context


logger = logging.getLogger(__name__)

_DEFAULT_IMAGE_WIDTH: Any = Inches(5.5)  # module-level singleton for ruff B008
# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _add_classification_header(doc: Document, classification: str = "CUI") -> None:
    """Add header/footer with classification marking to all sections."""
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.text = classification
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        run.font.bold = True

        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.text = f"{classification} | UIAO Program | Generated {datetime.now():%Y-%m-%d}"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.runs[0]
        fr.font.size = Pt(7)
        fr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Add a heading with consistent navy styling."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def _set_default_styles(doc: Document) -> None:
    """Set Calibri as the default font and tighten paragraph spacing document-wide."""
    from docx.oxml.ns import qn  # noqa: F401
    from docx.oxml import OxmlElement  # noqa: F401

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.15

    # Heading 1
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(4)

    # Heading 2
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(3)


def _add_narrative(doc: Document, text: Any) -> None:
    """Add narrative paragraph(s) with Calibri 11pt and clean spacing."""
    if not text:
        return
    for para_text in str(text).split("\n\n"):
        para_text = para_text.strip()
        if not para_text:
            continue
        p = doc.add_paragraph(para_text)
        p.style = doc.styles["Normal"]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.line_spacing = 1.15
        for run in p.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(11)


def _add_figure_caption(doc: Document, counter: list, title: str) -> None:
    """Add a 'Figure N: Title' caption below an image."""
    counter[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(14)
    label = cap.add_run(f"Figure {counter[0]}: ")
    label.font.name = "Calibri"
    label.font.size = Pt(9)
    label.font.bold = True
    label.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    body = cap.add_run(title)
    body.font.name = "Calibri"
    body.font.size = Pt(9)
    body.font.italic = True
    body.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_table_caption(doc: Document, counter: list, title: str) -> None:
    """Add a 'Table N: Title' caption above a table."""
    counter[0] += 1
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.space_before = Pt(14)
    cap.paragraph_format.space_after = Pt(3)
    label = cap.add_run(f"Table {counter[0]}: ")
    label.font.name = "Calibri"
    label.font.size = Pt(9)
    label.font.bold = True
    label.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    body = cap.add_run(title)
    body.font.name = "Calibri"
    body.font.size = Pt(9)
    body.font.italic = True
    body.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


def _add_image_safe(
    doc: Document,
    image_name: str,
    visuals_dir: Path,
    fig_counter: list | None = None,
    caption: str = "",
    width: Any = _DEFAULT_IMAGE_WIDTH,
) -> bool:
    """Add an image with optional Figure N caption below it."""
    img_path = visuals_dir / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if fig_counter is not None and caption:
            _add_figure_caption(doc, fig_counter, caption)
        return True
    logger.warning("Image not found: %s", img_path)
    return False


def _add_compliance_table(doc: Document, matrix: list[dict], tbl_counter: list | None = None) -> None:
    """Add the Unified Compliance Matrix as a formatted Word table."""
    if not matrix:
        return
    if tbl_counter is not None:
        _add_table_caption(doc, tbl_counter, "Unified Compliance & Maturity Matrix")
    headers = [
        "UIAO Pillar",
        "CISA ZT Pillar",
        "Target Maturity",
        "NIST 800-53 Controls",
        "Mission Impact",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Medium Shading 1 Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

        for entry in matrix:
            if not isinstance(entry, dict):
                continue
            row = table.add_row()
            row.cells[0].text = entry.get("pillar", "")
            row.cells[1].text = entry.get("cisa_pillar", "")
            row.cells[2].text = entry.get("target_maturity", "")
            controls = entry.get("nist_controls", [])
            row.cells[3].text = ", ".join(controls) if isinstance(controls, list) else str(controls)
            row.cells[4].text = entry.get("impact_statement", "")
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)


def _add_evidence_table(doc: Document, tbl_counter: list | None = None) -> None:
    """Add the FedRAMP Audit Evidence Summary table."""
    if tbl_counter is not None:
        _add_table_caption(doc, tbl_counter, "FedRAMP 20x Audit Evidence Summary — NIST 800-53 Rev 5 Control Mapping")
    evidence_map = [
        ("V1: Identity-to-IP Mapping", "U + A (The Gate)", "IA-2, AC-19, CM-8"),
        ("V2: INR Fabric", "O (The Network)", "AC-4"),
        ("V3: 20x Governance Loop", "Governance (The Hub)", "CA-7, IR-4"),
        ("V4: Modernization Atlas", "Strategy (The Journey)", "Program Vision / TIC 3.0"),
        ("V5: Cryptographic Trust Chain", "Security (The Lock)", "SC-8"),
    ]
    headers = ["Visual Title", "Architectural Pillar", "NIST Control(s)"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)
    for title, pillar, controls in evidence_map:
        row = table.add_row()
        row.cells[0].text = title
        row.cells[1].text = pillar
        row.cells[2].text = controls
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# Scratch builder
# ---------------------------------------------------------------------------


def _build_from_scratch(context: dict, visuals_dir: Path) -> Document:
    """Build a complete styled DOCX programmatically."""
    doc = Document()
    _set_default_styles(doc)
    fig = [0]  # mutable figure counter
    tbl = [0]  # mutable table counter
    lb = context.get("leadership_briefing", {})
    if not isinstance(lb, dict):
        lb = {}
    classification = context.get("classification", "CUI")
    _add_classification_header(doc, classification)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Unified Identity-Addressing-Overlay Architecture (UIAO)")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(f"Leadership Briefing - Version {context.get('version', '1.0')}")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = date_p.add_run(f"Generated: {datetime.now():%B %d, %Y}")
    dr.font.size = Pt(10)
    dr.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Disclaimer callout — rendered as amber notice box on title page
    disclaimer_text = context.get("disclaimer", "").strip()
    if disclaimer_text:
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OxmlElement

        d = doc.add_paragraph()
        d.paragraph_format.space_before = Pt(14)
        d.paragraph_format.space_after = Pt(14)
        d.paragraph_format.left_indent = Inches(0.3)
        d.paragraph_format.right_indent = Inches(0.3)
        d.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pPr = d._p.get_or_add_pPr()
        shd = _OxmlElement("w:shd")
        shd.set(_qn("w:val"), "clear")
        shd.set(_qn("w:color"), "auto")
        shd.set(_qn("w:fill"), "FFF3CD")
        pPr.append(shd)
        label_run = d.add_run("NOTICE: ")
        label_run.font.name = "Calibri"
        label_run.font.size = Pt(9)
        label_run.font.bold = True
        label_run.font.color.rgb = RGBColor(0x85, 0x6A, 0x04)
        text_run = d.add_run(disclaimer_text)
        text_run.font.name = "Calibri"
        text_run.font.size = Pt(9)
        text_run.font.color.rgb = RGBColor(0x66, 0x50, 0x00)

    doc.add_page_break()

    # TOC placeholder
    _add_heading(doc, "Table of Contents", level=1)
    toc_p = doc.add_paragraph()
    toc_p.add_run("[Table of Contents - right-click and select 'Update Field' in Word]")
    toc_p.runs[0].font.size = Pt(9)
    toc_p.runs[0].font.italic = True
    doc.add_page_break()

    # Sections
    _add_heading(doc, "Executive Summary")
    _add_narrative(doc, lb.get("executive_summary"))

    _add_heading(doc, "Program Overview")
    _add_narrative(doc, lb.get("program_overview"))

    _add_heading(doc, "Why Modernization Is Required")
    _add_narrative(doc, lb.get("modernization_need"))

    _add_heading(doc, "Program Vision")
    _add_narrative(doc, lb.get("program_vision"))

    # Five Control Planes
    _add_heading(doc, "The Five Control Planes")
    for i, plane in enumerate(lb.get("control_planes", []), 1):
        if isinstance(plane, dict):
            _add_heading(doc, f"{i}. {plane.get('name', '')}", level=2)
            _add_narrative(doc, plane.get("narrative"))

    # Eight Core Concepts
    _add_heading(doc, "Eight Core Concepts")
    for i, concept in enumerate(lb.get("core_concepts", []), 1):
        if isinstance(concept, dict):
            _add_heading(doc, f"{i}. {concept.get('name', '')}", level=2)
            _add_narrative(doc, concept.get("narrative"))

    _add_heading(doc, "Frozen State Analysis")
    _add_narrative(doc, lb.get("frozen_state"))

    _add_heading(doc, "Program Outcomes")
    _add_narrative(doc, lb.get("outcomes"))

    # Vibrant Visualizations
    doc.add_page_break()
    _add_heading(doc, "Vibrant Visualizations")
    _add_narrative(
        doc,
        "The following diagrams represent key architectural views of the UIAO program. "
        "Each figure is referenced by number in the architecture sections above.",
    )
    visuals = [
        (
            "Modernization Journey",
            "uiao-vibrant-modernization-atlas.png",
            "UIAO Modernization Atlas — transition from fragmented legacy to unified control planes",
        ),
        (
            "FedRAMP 20x Governance Loop",
            "uiao-vibrant-20x-governance-hub.png",
            "FedRAMP 20x Governance Hub — continuous compliance evidence loop",
        ),
        (
            "Identity-to-IP Architecture",
            "uiao-vibrant-u-plus-a-mapping.png",
            "Identity-to-IP Mapping — Entra ID as root namespace for all addressing decisions",
        ),
    ]
    for title_text, img, caption in visuals:
        _add_heading(doc, title_text, level=2)
        _add_image_safe(doc, img, visuals_dir, fig_counter=fig, caption=caption)

    # Maturity radar chart
    radar_path = visuals_dir / "dynamic-maturity-radar.png"
    if radar_path.exists():
        doc.add_page_break()
        _add_heading(doc, "CISA Zero Trust Maturity Assessment")
        _add_image_safe(
            doc,
            "dynamic-maturity-radar.png",
            visuals_dir,
            fig_counter=fig,
            caption="CISA Zero Trust Maturity Radar — current vs. target maturity across all five pillars",
        )

    # PlantUML-rendered architecture diagrams
    plantuml_dir = visuals_dir.parent / "plantuml"
    if plantuml_dir.is_dir():
        plantuml_pngs = sorted(plantuml_dir.glob("*.png"))
        if plantuml_pngs:
            doc.add_page_break()
            _add_heading(doc, "Architecture Diagrams (PlantUML)")
            _add_narrative(
                doc,
                "The following PlantUML diagrams are generated from the UIAO canon YAML. "
                "Each diagram is versioned alongside the architecture source and reflects "
                "the current state of the authoritative model.",
            )
            for png in plantuml_pngs:
                diagram_title = png.stem.replace("-", " ").replace("_", " ").title()
                _add_heading(doc, diagram_title, level=2)
                _add_image_safe(
                    doc,
                    png.name,
                    plantuml_dir,
                    fig_counter=fig,
                    caption=f"{diagram_title} — PlantUML architecture diagram",
                )

    # Gemini AI-generated visuals
    gemini_dir = visuals_dir.parent / "gemini"
    if gemini_dir.is_dir():
        gemini_pngs = sorted(gemini_dir.glob("*.png"))
        if gemini_pngs:
            doc.add_page_break()
            _add_heading(doc, "AI-Generated Visuals (Gemini)")
            for png in gemini_pngs:
                gemini_title = png.stem.replace("-", " ").replace("_", " ").title()
                _add_heading(doc, gemini_title, level=2)
                _add_image_safe(
                    doc,
                    png.name,
                    gemini_dir,
                    fig_counter=fig,
                    caption=f"{gemini_title} — AI-generated conceptual visualization",
                )

    # FedRAMP Evidence Summary
    doc.add_page_break()
    _add_heading(doc, "FedRAMP 20x Audit Evidence Summary")
    _add_narrative(
        doc,
        "Direct mapping of UIAO architecture to NIST 800-53 Rev 5 controls.",
    )
    _add_narrative(
        doc,
        f"Direct mapping of UIAO architecture to NIST 800-53 Rev 5 controls. "
        f"See Table {tbl[0] + 1} for the full evidence map.",
    )
    _add_evidence_table(doc, tbl_counter=tbl)

    # Compliance Matrix
    doc.add_page_break()
    _add_heading(doc, "Unified Compliance & Maturity Matrix")
    matrix = context.get("unified_compliance_matrix", [])
    if matrix:
        _add_compliance_table(doc, matrix, tbl_counter=tbl)
        p = doc.add_paragraph()
        p.add_run(
            "Auditor Note: All controls listed above are continuously "
            "monitored via the UIAO Governance Plane (V3) and reported "
            "through the ServiceNow SCuBA integration."
        ).font.italic = True

    # Figure & table index note
    if fig[0] > 0 or tbl[0] > 0:
        doc.add_page_break()
        _add_heading(doc, "List of Figures and Tables")
        summary = doc.add_paragraph()
        summary.add_run(
            f"This document contains {fig[0]} figure(s) and {tbl[0]} table(s). "
            "All figures are captioned with a sequential Figure number and descriptive title below the image. "
            "All tables are captioned with a sequential Table number and title above the table body."
        )
        summary.runs[0].font.size = Pt(10)
        summary.runs[0].font.italic = True

    return doc


# ---------------------------------------------------------------------------
# Markdown-to-DOCX topic document builder  (Issues 1 & 2)
# ---------------------------------------------------------------------------

_MD_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_MD_HR_RE = re.compile(r"^-{3,}$")
_MD_FRONTMATTER_RE = re.compile(r"^---\n.*?^---\n", re.DOTALL | re.MULTILINE)


_IMG_TAG_RE = re.compile(r'<img\\s+[^>]*src="([^"]+)"[^>]*>', re.IGNORECASE)
_MD_TABLE_ROW_RE = re.compile(r"^\\|(.+)\\|\\s*$")
_MD_TABLE_SEP_RE = re.compile(r"^[\\s|:\\-]+$")


def _md_to_docx(doc: Document, md_text: str) -> None:
    """Render simplified Markdown into an existing Document.

    Handles H1-H3 headings, paragraphs, horizontal rules, pipe tables,
    and <img> tags (resolved to embedded pictures).
    Inline bold (**text**) and italic (*text*) are also supported.
    """
    md_text = _MD_FRONTMATTER_RE.sub("", md_text).strip()

    lines = md_text.splitlines()
    i = 0
    para_buffer: list[str] = []

    def _flush_paragraph() -> None:
        nonlocal para_buffer
        if not para_buffer:
            return
        text = " ".join(para_buffer).strip()
        para_buffer = []
        if not text:
            return
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        remaining = text
        while remaining:
            bold_m = re.search(r"\\*\\*(.+?)\\*\\*", remaining)
            ital_m = re.search(r"(?<!\\*)\\*(?!\\*)(.+?)(?<!\\*)\\*(?!\\*)", remaining)
            first = None
            if bold_m and ital_m:
                first = bold_m if bold_m.start() <= ital_m.start() else ital_m
            elif bold_m:
                first = bold_m
            elif ital_m:
                first = ital_m
            if first is None:
                r = p.add_run(remaining)
                r.font.name = "Calibri"
                r.font.size = Pt(11)
                break
            if first.start() > 0:
                r = p.add_run(remaining[: first.start()])
                r.font.name = "Calibri"
                r.font.size = Pt(11)
            r = p.add_run(first.group(1))
            r.font.name = "Calibri"
            r.font.size = Pt(11)
            r.font.bold = first is bold_m
            r.font.italic = first is not bold_m
            remaining = remaining[first.end() :]

    while i < len(lines):
        line = lines[i].rstrip()

        if _MD_HR_RE.match(line):
            _flush_paragraph()
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        m = _MD_HEADING_RE.match(line)
        if m:
            _flush_paragraph()
            level = len(m.group(1))
            text = re.sub(r"\\*+(.+?)\\*+", r"\\1", m.group(2))
            h = doc.add_heading(text, level=min(level, 3))
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
            i += 1
            continue

        img_m = _IMG_TAG_RE.search(line)
        if img_m:
            _flush_paragraph()
            src = img_m.group(1)
            img_path = Path(src)
            if img_path.exists():
                try:
                    doc.add_picture(str(img_path), width=Inches(6.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass
            i += 1
            continue

        if _MD_TABLE_ROW_RE.match(line):
            _flush_paragraph()
            table_rows: list[list[str]] = []
            while i < len(lines) and _MD_TABLE_ROW_RE.match(lines[i].rstrip()):
                row_text = lines[i].strip().strip("|")
                cells = [c.strip() for c in row_text.split("|")]
                inner = lines[i].strip().replace("|", "").replace(" ", "").replace("-", "").replace(":", "")
                if inner:
                    table_rows.append(cells)
                i += 1
            if table_rows:
                num_cols = max(len(r) for r in table_rows)
                tbl = doc.add_table(rows=0, cols=num_cols)
                tbl.style = "Light List Accent 1"
                tbl.autofit = True
                for ri, row_cells in enumerate(table_rows):
                    row = tbl.add_row()
                    for ci in range(num_cols):
                        cell_text = row_cells[ci] if ci < len(row_cells) else ""
                        row.cells[ci].text = cell_text
                        for p in row.cells[ci].paragraphs:
                            for run in p.runs:
                                run.font.size = Pt(9)
                                if ri == 0:
                                    run.font.bold = True
            continue

        if not line:
            _flush_paragraph()
            i += 1
            continue

        para_buffer.append(line)
        i += 1

    _flush_paragraph()


def build_topic_docx(
    md_path: Path,
    exports_dir: Path,
    classification: str = "CUI",
) -> Path:
    """Convert a single Markdown document into a styled DOCX.

    Args:
        md_path: Path to the source Markdown file (from docs/).
        exports_dir: Root exports directory; output goes to exports_dir/docx/.
        classification: Header/footer marking (default: CUI).

    Returns:
        Path to the generated DOCX file.
    """
    doc = Document()
    _set_default_styles(doc)
    # Apply 1" margins per canonical style guide (margin-safe, no fixed widths)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    _add_classification_header(doc, classification)
    md_text = md_path.read_text(encoding="utf-8")
    _md_to_docx(doc, md_text)
    docx_dir = exports_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    out_path = docx_dir / (md_path.stem + ".docx")
    doc.save(str(out_path))
    logger.info("Topic DOCX exported -> %s", out_path)
    return out_path


def generate_all_topic_docs(
    docs_dir: Path | None = None,
    exports_dir: Path | None = None,
    classification: str = "CUI",
) -> list[Path]:
    """Convert all Markdown documents in docs_dir to DOCX.

    Skips README, CHANGELOG, MANIFEST, index files. The rich Leadership
    Briefing DOCX is generated separately by build_rich_docx().

    Returns:
        List of generated DOCX Path objects.
    """
    settings = get_settings()
    if docs_dir is None:
        docs_dir = settings.root_dir / "docs"
    if exports_dir is None:
        exports_dir = settings.exports_dir
    docs_dir = Path(docs_dir)
    if not docs_dir.exists():
        logger.warning("docs_dir not found: %s", docs_dir)
        return []
    _SKIP = {"readme", "changelog", "manifest", "index"}
    # Issue 3: retire the legacy lowercase leadership_briefing_v1.0.docx
    # which was generated by an older pipeline. The authoritative file is
    # UIAO_Leadership_Briefing_v1.0.docx produced by build_rich_docx().
    _LEGACY_FILES = {"leadership_briefing_v1.0.docx"}
    docx_dir = Path(exports_dir) / "docx"
    for legacy in _LEGACY_FILES:
        legacy_path = docx_dir / legacy
        if legacy_path.exists():
            legacy_path.unlink()
            logger.info("Removed stale legacy file: %s", legacy_path)

    generated: list[Path] = []
    for md_path in sorted(docs_dir.glob("*.md")):
        if md_path.stem.lower() in _SKIP:
            continue
        # Skip the leadership_briefing Markdown — its DOCX is build_rich_docx()
        if "leadership_briefing" in md_path.stem.lower():
            continue
        try:
            out = build_topic_docx(md_path, exports_dir, classification=classification)
            generated.append(out)
        except Exception as exc:  # noqa: BLE001
            logger.warning("Failed to convert %s: %s", md_path.name, exc)
    return generated


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------


def build_rich_docx(
    *,
    canon_path: Path | None = None,
    data_dir: Path | None = None,
    templates_dir: Path | None = None,
    visuals_dir: Path | None = None,
    exports_dir: Path | None = None,
) -> Path:
    """Generate a rich DOCX leadership briefing and return the output path.

    All directory arguments default to values derived from ``Settings``.
    """
    settings = get_settings()
    canon_path = canon_path or settings.canon_dir / "uiao_leadership_briefing_v1.0.yaml"
    data_dir = data_dir or settings.data_dir
    templates_dir = templates_dir or settings.templates_dir
    visuals_dir = visuals_dir or settings.visuals_dir
    exports_dir = exports_dir or settings.exports_dir

    logger.info("Loading UIAO context...")
    context = load_context(canon_path=canon_path, data_dir=data_dir)

    docx_dir = exports_dir / "docx"
    docx_dir.mkdir(parents=True, exist_ok=True)
    out_path = docx_dir / "UIAO_Leadership_Briefing_v1.0.docx"

    # Try template-based approach first
    tpl_path = templates_dir / "leadership_briefing_v1.0.docx"
    if HAS_DOCXTPL and tpl_path.exists():
        logger.info("Using docxtpl template: %s", tpl_path)
        tpl = DocxTemplate(str(tpl_path))
        context["today"] = datetime.now().strftime("%B %d, %Y")
        context["compliance_table"] = context.get("unified_compliance_matrix", [])
        tpl.render(context)
        tpl.save(str(out_path))
    else:
        logger.info("Building styled DOCX from scratch...")
        doc = _build_from_scratch(context, visuals_dir)
        doc.save(str(out_path))

    logger.info("Rich DOCX exported -> %s", out_path)
    return out_path
